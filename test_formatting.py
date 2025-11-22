#!/usr/bin/env python3
"""Enhanced test script to check formatting improvements in generated DOCX."""

from pathlib import Path
from docx import Document
from docx.shared import Inches

def analyze_docx_formatting(docx_path: Path):
    """Analyze formatting in a DOCX file with focus on colors and column widths."""
    doc = Document(docx_path)
    
    print(f"📄 Analyzing: {docx_path}")
    print(f"📊 Paragraphs: {len(doc.paragraphs)}")
    print(f"📋 Tables: {len(doc.tables)}")
    print()
    
    # Check paragraphs for colors and formatting
    color_found = False
    for i, para in enumerate(doc.paragraphs[:15]):  # Check more paragraphs
        if para.text.strip():
            print(f"Para {i+1}: {para.text[:80]}...")
            
            # Check runs for formatting and colors
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    formatting = []
                    if run.bold:
                        formatting.append("BOLD")
                    if run.italic:
                        formatting.append("ITALIC")
                    
                    # Check for colors
                    if hasattr(run.font, 'color') and run.font.color.rgb:
                        color = run.font.color.rgb
                        try:
                            # Try to access RGB values
                            r, g, b = color
                            color_str = f"RGB({r},{g},{b})"
                            formatting.append(f"COLOR:{color_str}")
                            if g > r and g > b:  # Greenish
                                formatting.append("🟢GREEN!")
                                color_found = True
                        except (TypeError, ValueError):
                            formatting.append("COLOR:DETECTED")
                            color_found = True
                    
                    if formatting:
                        print(f"  Run {j+1}: '{run.text[:30]}' → {' | '.join(formatting)}")
            
            # Check alignment
            if para.alignment:
                print(f"  Alignment: {para.alignment}")
            print()
    
    if not color_found:
        print("⚠️  No green colors detected in paragraphs")
    
    # Check tables with detailed column width analysis
    for i, table in enumerate(doc.tables):
        print(f"🗃️  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        
        # Check column widths
        print("   Column widths:")
        for col_idx, column in enumerate(table.columns):
            width_inches = column.width.inches if column.width else "auto"
            print(f"     Col {col_idx+1}: {width_inches} inches")
        
        # Check header row
        if table.rows:
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            print(f"   Headers: {' | '.join(headers[:5])}")
        
        # Check first few data rows
        for row_idx, row in enumerate(table.rows[1:4]):  # Skip header, show 3 data rows
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text[:15])
            if row_data:
                print(f"   Row {row_idx+2}: {' | '.join(row_data)}")
        print()

if __name__ == "__main__":
    docx_file = Path("tmp/387-test.docx")
    if docx_file.exists():
        analyze_docx_formatting(docx_file)
    else:
        print(f"❌ File not found: {docx_file}")