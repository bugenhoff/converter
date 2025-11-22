#!/usr/bin/env python3
"""Quick test script to check the formatting in generated DOCX."""

from pathlib import Path
from docx import Document

def analyze_docx_formatting(docx_path: Path):
    """Analyze formatting in a DOCX file."""
    doc = Document(docx_path)
    
    print(f"📄 Analyzing: {docx_path}")
    print(f"📊 Paragraphs: {len(doc.paragraphs)}")
    print(f"📋 Tables: {len(doc.tables)}")
    print()
    
    # Check paragraphs
    for i, para in enumerate(doc.paragraphs[:10]):  # First 10 paragraphs
        if para.text.strip():
            print(f"Para {i+1}: {para.text[:80]}...")
            
            # Check runs for formatting
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    formatting = []
                    if run.bold:
                        formatting.append("BOLD")
                    if run.italic:
                        formatting.append("ITALIC")
                    if hasattr(run.font, 'color') and run.font.color.rgb:
                        color = run.font.color.rgb
                        formatting.append(f"COLOR({color.r},{color.g},{color.b})")
                    
                    if formatting:
                        print(f"  Run {j+1}: '{run.text[:30]}' → {' | '.join(formatting)}")
            
            # Check alignment
            if para.alignment:
                print(f"  Alignment: {para.alignment}")
            print()
    
    # Check tables
    for i, table in enumerate(doc.tables):
        print(f"🗃️  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        
        # Check first few cells for content
        for row_idx, row in enumerate(table.rows[:3]):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text[:20])
            if row_data:
                print(f"   Row {row_idx+1}: {' | '.join(row_data)}")
        print()

if __name__ == "__main__":
    docx_file = Path("tmp/387-test.docx")
    if docx_file.exists():
        analyze_docx_formatting(docx_file)
    else:
        print(f"❌ File not found: {docx_file}")