"""Uzbek Latin -> Cyrillic transliteration helpers for DOCX content."""

from __future__ import annotations

import io
import re
from typing import Iterable

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency in some environments
    Document = None


_APOSTROPHE_PATTERN = re.compile(r"[ʻʼ’‘`´ʹʽ‛]")
_MULTI_CHAR_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"o'", re.IGNORECASE), "ў"),
    (re.compile(r"g'", re.IGNORECASE), "ғ"),
    (re.compile(r"sh", re.IGNORECASE), "ш"),
    (re.compile(r"ch", re.IGNORECASE), "ч"),
    (re.compile(r"ng", re.IGNORECASE), "нг"),
    (re.compile(r"ya", re.IGNORECASE), "я"),
    (re.compile(r"yo", re.IGNORECASE), "ё"),
    (re.compile(r"yu", re.IGNORECASE), "ю"),
    (re.compile(r"ye", re.IGNORECASE), "е"),
]
_SINGLE_CHAR_MAP: dict[str, str] = {
    "a": "а",
    "b": "б",
    "c": "ц",
    "d": "д",
    "e": "е",
    "f": "ф",
    "g": "г",
    "h": "ҳ",
    "i": "и",
    "j": "ж",
    "k": "к",
    "l": "л",
    "m": "м",
    "n": "н",
    "o": "о",
    "p": "п",
    "q": "қ",
    "r": "р",
    "s": "с",
    "t": "т",
    "u": "у",
    "v": "в",
    "w": "в",
    "x": "х",
    "y": "й",
    "z": "з",
}


def transliterate_uz_latin_to_cyr(text: str) -> str:
    """Transliterate Uzbek text from Latin to Cyrillic."""
    if not text:
        return text

    transliterated = _APOSTROPHE_PATTERN.sub("'", text)

    for pattern, replacement in _MULTI_CHAR_PATTERNS:
        transliterated = pattern.sub(
            lambda match: _apply_case(match.group(0), replacement),
            transliterated,
        )

    output_chars: list[str] = []
    for char in transliterated:
        lower_char = char.lower()
        mapped = _SINGLE_CHAR_MAP.get(lower_char)
        if not mapped:
            output_chars.append(char)
            continue

        output_chars.append(mapped.upper() if char.isupper() else mapped)

    return "".join(output_chars)


def transliterate_docx_bytes(docx_bytes: bytes) -> bytes:
    """Transliterate paragraphs, tables and headers/footers in DOCX bytes."""
    if Document is None:
        raise ImportError("python-docx is required for DOCX transliteration")

    doc = Document(io.BytesIO(docx_bytes))

    _transliterate_paragraphs(doc.paragraphs)
    _transliterate_tables(doc.tables)

    for section in doc.sections:
        for header_footer in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            _transliterate_paragraphs(header_footer.paragraphs)
            _transliterate_tables(header_footer.tables)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _apply_case(source: str, replacement: str) -> str:
    if source.isupper():
        return replacement.upper()
    if len(source) > 1 and source[0].isupper() and source[1:].islower():
        if len(replacement) == 1:
            return replacement.upper()
        return replacement[0].upper() + replacement[1:]
    return replacement


def _transliterate_paragraphs(paragraphs: Iterable) -> None:
    for paragraph in paragraphs:
        if paragraph.runs:
            for run in paragraph.runs:
                if run.text:
                    run.text = transliterate_uz_latin_to_cyr(run.text)
            continue
        if paragraph.text:
            paragraph.text = transliterate_uz_latin_to_cyr(paragraph.text)


def _transliterate_tables(tables: Iterable) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _transliterate_paragraphs(cell.paragraphs)
                _transliterate_tables(cell.tables)
