"""Groq LLM-based PDF to DOCX conversion with vision models."""

from __future__ import annotations

import base64
import logging
import subprocess
import tempfile
from collections import deque
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, TypeVar

try:
    import groq
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from PIL import Image
    from pdf2image import convert_from_bytes
except ImportError:
    groq = None
    Document = None
    convert_from_bytes = None

from ..config.settings import settings

logger = logging.getLogger(__name__)
T = TypeVar("T")


class GroqConversionError(RuntimeError):
    """Signals that Groq LLM conversion failed."""


@dataclass
class _GroqTask:
    page_items: list[tuple[int, T]]
    image_max_side: int
    retries_left: int


def convert_pdf_to_docx_via_groq(source_path: Path, output_dir: Path) -> Path:
    """Convert PDF to DOCX using Groq vision LLM."""
    
    if not groq or not Document:
        raise ImportError("groq and python-docx are required for LLM conversion")
    
    if not settings.groq_api_key:
        raise GroqConversionError("GROQ_API_KEY is not configured")

    source_path = Path(source_path)
    if not source_path.exists():
        raise FileNotFoundError(f"Source file {source_path} was not found")

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{source_path.stem}.docx"

    logger.info("Converting PDF to images for LLM processing: %s", source_path)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        # Convert PDF to images
        images = _pdf_to_images(source_path, Path(temp_dir))
        
        if not images:
            raise GroqConversionError("Failed to extract images from PDF")
        
        logger.info("Extracted %d page images, sending to Groq LLM", len(images))
        
        # Process with Groq LLM
        structured_content = _process_images_with_groq(images)
        
        # Create DOCX from structured content  
        _create_docx_from_content(structured_content, docx_path)
        
        logger.info("LLM-based conversion completed: %s", docx_path)

    if not docx_path.exists():
        raise GroqConversionError("Output file is missing after conversion")
    
    if docx_path.stat().st_size == 0:
        raise GroqConversionError("Conversion produced an empty DOCX file")

    return docx_path


def convert_pdf_bytes_to_docx_via_groq(pdf_bytes: bytes, original_name: str) -> bytes:
    """Convert PDF bytes to DOCX bytes using Groq vision LLM (in-memory processing)."""
    
    if not groq or not Document or not convert_from_bytes:
        raise ImportError("groq, python-docx, and pdf2image are required for memory conversion")
    
    if not settings.groq_api_key:
        raise GroqConversionError("GROQ_API_KEY is not configured")
    
    logger.info("Converting PDF bytes to DOCX in memory: %s (%d bytes)", original_name, len(pdf_bytes))
    logger.info(
        "Groq config: model=%s max_tokens=%d batch_size=%d max_side=%d dpi=%d",
        settings.groq_model,
        settings.groq_max_tokens,
        settings.groq_batch_size,
        settings.groq_image_max_side,
        settings.groq_pdf_image_dpi,
    )
    
    try:
        # Convert PDF bytes to images in memory
        images = _pdf_bytes_to_images(pdf_bytes)
        
        if not images:
            raise GroqConversionError("Failed to extract images from PDF bytes")
        
        logger.info("Extracted %d page images from memory, sending to Groq LLM", len(images))
        
        # Process with Groq LLM
        structured_content = _process_pil_images_with_groq(images)
        
        # Create DOCX in memory
        docx_bytes = _create_docx_bytes_from_content(structured_content)
        
        logger.info("Memory-based LLM conversion completed for %s: %d bytes", 
                   original_name, len(docx_bytes))
        
        return docx_bytes
        
    except Exception as e:
        raise GroqConversionError(f"Memory conversion failed for {original_name}: {e}") from e


def _pdf_bytes_to_images(pdf_bytes: bytes) -> list[Image.Image]:
    """Convert PDF bytes to PIL Images in memory."""
    try:
        # Use pdf2image to convert bytes directly to PIL Images
        images = convert_from_bytes(
            pdf_bytes,
            dpi=settings.groq_pdf_image_dpi,
            fmt='PNG',
            use_pdftocairo=True  # Better quality rendering
        )
        logger.debug(
            "Converted PDF bytes to %d PIL images (dpi=%d)",
            len(images),
            settings.groq_pdf_image_dpi,
        )
        return images
        
    except Exception as e:
        raise GroqConversionError(f"Failed to convert PDF bytes to images: {e}") from e


def _process_pil_images_with_groq(images: list[Image.Image]) -> dict[str, Any]:
    """Process PIL Images with simple fixed-size batching (no adaptive logic)."""
    if not images:
        raise GroqConversionError("No images to process")

    client = groq.Groq(api_key=settings.groq_api_key)
    batch_size = max(1, settings.groq_batch_size)
    all_results: list[dict[str, Any]] = []

    logger.info(
        "Simple Groq scheduler started: pages=%d batch_size=%d",
        len(images),
        batch_size,
    )

    for batch_num, start in enumerate(range(0, len(images), batch_size)):
        batch_images = images[start : start + batch_size]
        start_page = start + 1
        end_page = start + len(batch_images)

        logger.info(
            "Groq request %d/%d: pages=%d-%d count=%d",
            batch_num + 1,
            -(-len(images) // batch_size),  # ceil division
            start_page,
            end_page,
            len(batch_images),
        )

        try:
            result = _process_pil_batch_with_groq(
                client,
                batch_images,
                start_page,
                batch_num,
                settings.groq_image_max_side,
            )
            all_results.append(result)
        except Exception as exc:
            raise GroqConversionError(
                f"Groq request failed for pages {start_page}-{end_page}: {exc}"
            ) from exc

    merged = _merge_batch_results(all_results)
    logger.info(
        "Simple Groq scheduler completed: requests=%d final_pages=%d",
        len(all_results),
        len(merged.get("pages", [])),
    )
    return merged


def _process_images_with_adaptive_strategy(
    images: list[T],
    processor: Callable[[Any, list[T], int, int, int], dict[str, Any]],
) -> dict[str, Any]:
    """Process pages with adaptive split/downscale/retry strategy."""

    if not images:
        raise GroqConversionError("No images to process")

    client = groq.Groq(api_key=settings.groq_api_key)
    tasks = _build_initial_tasks(images)

    all_results: list[dict[str, Any]] = []
    max_requests = settings.groq_max_requests_per_document
    requests_made = 0

    logger.info(
        "Adaptive Groq scheduler started: pages=%d initial_batch=%d min_batch=%d max_side=%d min_side=%d max_requests=%s retries=%d",
        len(images),
        settings.groq_batch_size,
        settings.groq_min_batch_size,
        settings.groq_image_max_side,
        settings.groq_min_image_max_side,
        "unlimited" if max_requests == 0 else max_requests,
        settings.groq_retry_per_task,
    )

    while tasks:
        task = tasks.popleft()
        expected_pages = [page_number for page_number, _ in task.page_items]

        if max_requests > 0 and requests_made >= max_requests:
            raise GroqConversionError(
                f"Groq request limit reached ({requests_made}/{max_requests}) before full document coverage"
            )

        requests_made += 1
        request_cap = "unlimited" if max_requests == 0 else str(max_requests)
        logger.info(
            "Groq request %d/%s: pages=%s pages_in_task=%d image_max_side=%d max_tokens=%d",
            requests_made,
            request_cap,
            _format_page_numbers(expected_pages),
            len(expected_pages),
            task.image_max_side,
            settings.groq_max_tokens,
        )

        try:
            page_payload = [item for _, item in task.page_items]
            batch_result = processor(
                client,
                page_payload,
                expected_pages[0],
                requests_made - 1,
                task.image_max_side,
            )
            _validate_batch_pages(batch_result, expected_pages)
            all_results.append(batch_result)
            continue
        except Exception as exc:
            logger.warning(
                "Groq task failed for pages=%s image_max_side=%d retries_left=%d: %s",
                _format_page_numbers(expected_pages),
                task.image_max_side,
                task.retries_left,
                exc,
            )

        split_tasks = _split_task(task)
        if split_tasks:
            logger.info(
                "Adaptive action=split pages=%s into %s + %s",
                _format_page_numbers(expected_pages),
                _format_page_numbers([n for n, _ in split_tasks[0].page_items]),
                _format_page_numbers([n for n, _ in split_tasks[1].page_items]),
            )
            tasks.appendleft(split_tasks[1])
            tasks.appendleft(split_tasks[0])
            continue

        downscaled = _downscale_task(task)
        if downscaled:
            logger.info(
                "Adaptive action=downscale pages=%s old_side=%d new_side=%d",
                _format_page_numbers(expected_pages),
                task.image_max_side,
                downscaled.image_max_side,
            )
            tasks.appendleft(downscaled)
            continue

        if task.retries_left > 0:
            retried_task = _GroqTask(
                page_items=task.page_items,
                image_max_side=task.image_max_side,
                retries_left=task.retries_left - 1,
            )
            logger.info(
                "Adaptive action=retry pages=%s retries_left=%d",
                _format_page_numbers(expected_pages),
                retried_task.retries_left,
            )
            tasks.appendleft(retried_task)
            continue

        raise GroqConversionError(
            "Groq conversion failed after exhausting split/downscale/retry policy for pages "
            f"{_format_page_numbers(expected_pages)}"
        )

    if not all_results:
        raise GroqConversionError("All Groq tasks failed to process")

    merged = _merge_batch_results(all_results)
    logger.info(
        "Adaptive Groq scheduler completed: requests_made=%d final_pages=%d",
        requests_made,
        len(merged.get("pages", [])),
    )
    return merged


def _process_pil_batch_with_groq(
    client,
    pil_images: list[Image.Image],
    start_page: int,
    batch_idx: int,
    image_max_side: int,
) -> dict[str, Any]:
    """Process a batch of PIL Images with Groq LLM."""
    
    image_content = []
    for img in pil_images:
        # Resize image if too large to save tokens
        if max(img.size) > image_max_side:
            ratio = image_max_side / max(img.size)
            new_size = tuple(int(dim * ratio) for dim in img.size)
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        # Convert PIL Image to base64 bytes
        import io
        img_buffer = io.BytesIO()
        img.save(img_buffer, format='PNG', optimize=True)
        img_bytes = img_buffer.getvalue()
        b64_image = base64.b64encode(img_bytes).decode()
        
        image_content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{b64_image}"
            }
        })
    
    # Use existing prompt from the main function
    prompt = f"""Analyze these {len(pil_images)} pages (pages {start_page}-{start_page + len(pil_images) - 1}) and extract ALL content with EXACT formatting.

    JSON format:
    {{
        "title": "document title (batch 1 only)",
        "pages": [
            {{
                "page_number": {start_page},
                "sections": [
                    {{
                        "type": "header_row|heading|paragraph|table|list",
                        "content": "COMPLETE TEXT",
                        "formatting": {{
                            "bold": true/false,
                            "color": "green|blue|red|black",
                            "alignment": "left|center|right|justified",
                            "font_size": "small|normal|large"
                        }},
                        "layout": {{
                            "left_text": "text on left side",
                            "right_text": "text on right side",
                            "position": "header|body|footer"
                        }},
                        "table_data": {{
                            "headers": ["col1", "col2", "col3"],
                            "rows": [["cell1", "cell2", "cell3"]],
                            "has_borders": true,
                            "header_styling": true
                        }}
                    }}
                ]
            }}
        ]
    }}
    
    CRITICAL ANALYSIS:
    - COLORS: Identify text colors (green BUYRUGI, etc.)
    - BOLD: Mark all bold/emphasized text
    - TABLES: Extract complete table structure with borders
    - LAYOUT: Note date/number positioning (left vs right)
    - NUMBERS: Highlight numeric formatting
    - Extract EVERYTHING - no truncation
    """

    try:
        response = client.chat.completions.create(
            model=settings.groq_model,
            messages=[
                {
                    "role": "user", 
                    "content": [
                        {"type": "text", "text": prompt},
                        *image_content
                    ]
                }
            ],
            response_format={"type": "json_object"},
            max_tokens=settings.groq_max_tokens,
            temperature=0.1
        )
        
        content = response.choices[0].message.content
        logger.info("Raw Groq response for batch %d: %s", batch_idx + 1, content[:500] + "...")
        
        import json
        result = json.loads(content)
        return result
        
    except Exception as e:
        raise GroqConversionError(f"Groq API request failed: {e}") from e


def _create_docx_bytes_from_content(content: dict[str, Any]) -> bytes:
    """Create DOCX document from structured content and return as bytes."""
    import io
    
    # Create document in memory  
    doc = Document()
    
    # Use existing _create_docx_from_content logic but write to BytesIO
    def safe_text(value: Any) -> str:
        if isinstance(value, list):
            return ' '.join(str(item) for item in value if item)
        return str(value) if value else ""
    
    # Title comes from page sections — no separate heading to avoid duplication
    pages = content.get("pages", [])
    logger.info("Processing %d pages for DOCX creation", len(pages))
    
    for page_idx, page in enumerate(pages):
        page_num = page.get("page_number", page_idx + 1)
        sections = page.get("sections", [])
        logger.info("Page %d: processing %d sections", page_num, len(sections))
        
        for section_idx, section in enumerate(sections):
            section_type = section.get("type", "paragraph")
            text_content = safe_text(section.get("content", "")).strip()
            formatting = section.get("formatting", {})
            layout = section.get("layout", {})
            table_data = section.get("table_data", {})
            
            if not text_content and not table_data:
                if section_type == "table":
                    logger.warning("Skipping table section with empty table_data on page %d (section %d)", page_num, section_idx)
                else:
                    logger.debug("Skipping empty section %d on page %d", section_idx, page_num)
                continue
            
            logger.debug("Adding %s: %s", section_type, text_content[:100])
            
            # Handle different section types using existing functions
            if section_type == "header_row":
                _add_header_row(doc, layout, formatting, text_content)
            elif section_type == "heading":
                _add_formatted_heading(doc, text_content, formatting)
            elif section_type == "table" and table_data:
                _add_formatted_table(doc, table_data, text_content)
            elif section_type == "list":
                _add_formatted_list(doc, text_content, formatting)
            else:  # paragraph
                _add_formatted_paragraph(doc, text_content, formatting)
    
    # Save document to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    
    logger.info("DOCX document created in memory: %d bytes", len(docx_bytes))
    return docx_bytes


def _pdf_to_images(pdf_path: Path, output_dir: Path) -> list[Path]:
    """Convert PDF pages to PNG images."""
    try:
        # Use configurable DPI for image extraction quality
        subprocess.run([
            "pdftoppm", "-png", "-r", str(settings.groq_pdf_image_dpi),
            str(pdf_path), str(output_dir / "page")
        ], check=True, capture_output=True)
        
        # Sort files numerically: page-1.png, page-2.png, ...
        images = sorted(
            output_dir.glob("page-*.png"),
            key=lambda p: int(p.stem.split('-')[-1])
        )
        return images
        
    except subprocess.CalledProcessError as e:
        raise GroqConversionError(f"Failed to convert PDF to images: {e.stderr}") from e


def _process_images_with_groq(image_paths: list[Path]) -> dict[str, Any]:
    """Process images with adaptive Groq batching strategy."""

    merged = _process_images_with_adaptive_strategy(
        image_paths,
        _process_batch_with_groq,
    )
    _validate_merged_pages(merged, expected_pages=len(image_paths))
    return merged


def _process_batch_with_groq(
    client,
    image_paths: list[Path],
    start_page: int,
    batch_idx: int,
    image_max_side: int,
) -> dict[str, Any]:
    """Process a single batch of images (up to 5) with Groq LLM."""

    image_content = []
    for img_path in image_paths:
        # Resize image if too large to save tokens (reduced resolution)
        with Image.open(img_path) as img:
            if max(img.size) > image_max_side:
                ratio = image_max_side / max(img.size)
                new_size = tuple(int(dim * ratio) for dim in img.size)
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Save resized image
                resized_path = img_path.with_suffix('.resized.png')
                img.save(resized_path, 'PNG')
                img_path = resized_path
        
        # Encode to base64
        with open(img_path, 'rb') as f:
            b64_image = base64.b64encode(f.read()).decode()
        
        image_content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{b64_image}"
            }
        })

    # Advanced prompt for formatting and structure recognition with enhanced color detection
    prompt = f"""Analyze these {len(image_paths)} pages (pages {start_page}-{start_page + len(image_paths) - 1}) and extract ALL content with EXACT formatting and colors.

    JSON format:
    {{
        "title": "document title (batch 1 only)",
        "pages": [
            {{
                "page_number": {start_page},
                "sections": [
                    {{
                        "type": "header_row|heading|paragraph|table|list",
                        "content": "COMPLETE TEXT",
                        "formatting": {{
                            "bold": true/false,
                            "color": "green|blue|red|black|dark_blue|orange",
                            "alignment": "left|center|right|justified",
                            "font_size": "small|normal|large"
                        }},
                        "layout": {{
                            "left_text": "text on left side",
                            "right_text": "text on right side",
                            "position": "header|body|footer"
                        }},
                        "table_data": {{
                            "headers": ["col1", "col2", "col3"],
                            "rows": [["cell1", "cell2", "cell3"]],
                            "has_borders": true,
                            "header_styling": true,
                            "column_types": ["number", "text", "number", "number", "number"]
                        }}
                    }}
                ]
            }}
        ]
    }}
    
    CRITICAL COLOR ANALYSIS:
    - LOOK FOR ANY COLORED TEXT
    - BOLD TEXT: Mark all emphasized/bold text accurately  
    - TABLE STRUCTURE: Extract complete table with proper headers and data
    - COLUMN ANALYSIS: Identify column types (number/text) for width optimization
    - LAYOUT POSITIONING: Note left/right text alignment (dates vs numbers)
    - NUMBERS: Highlight all numeric content formatting
    
    SPECIFIC EXAMPLES TO LOOK FOR:
    - "BUYRUGI" text in GREEN color (this is very important!)
    - Date numbers (20, 25, 18, 387) should be marked as bold
    - Table columns: № (narrow numbers), Nomi (wide text), Soni/Narxi/Jami (medium numbers)
    - Any text that appears in colors other than black
    
    
    Extract EVERYTHING - no truncation allowed!
    """

    try:
        response = client.chat.completions.create(
            model=settings.groq_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        *image_content
                    ]
                }
            ],
            response_format={"type": "json_object"},
            max_tokens=settings.groq_max_tokens,
            temperature=0.1
        )
        
        content = response.choices[0].message.content
        logger.info("Raw Groq response for batch %d: %s", batch_idx + 1, content[:2000] + "...")
        
        import json
        result = json.loads(content)
        pages = result.get('pages', [])
        logger.info("Parsed JSON structure: title=%s, pages_count=%d, page_numbers=%s", 
                   type(result.get('title')), len(pages), 
                   [p.get('page_number') for p in pages])
        return result
        
    except Exception as e:
        raise GroqConversionError(f"Groq API request failed: {e}") from e


def _build_initial_tasks(images: list[T]) -> deque[_GroqTask]:
    batch_size = settings.groq_batch_size
    enumerated = [(idx + 1, image) for idx, image in enumerate(images)]
    tasks: deque[_GroqTask] = deque()
    for i in range(0, len(enumerated), batch_size):
        tasks.append(
            _GroqTask(
                page_items=enumerated[i:i + batch_size],
                image_max_side=settings.groq_image_max_side,
                retries_left=settings.groq_retry_per_task,
            )
        )
    return tasks


def _split_task(task: _GroqTask) -> tuple[_GroqTask, _GroqTask] | None:
    if len(task.page_items) <= settings.groq_min_batch_size:
        return None

    midpoint = len(task.page_items) // 2
    if midpoint < settings.groq_min_batch_size:
        return None

    left_pages = task.page_items[:midpoint]
    right_pages = task.page_items[midpoint:]
    if len(left_pages) < settings.groq_min_batch_size or len(right_pages) < settings.groq_min_batch_size:
        return None

    return (
        _GroqTask(
            page_items=left_pages,
            image_max_side=task.image_max_side,
            retries_left=task.retries_left,
        ),
        _GroqTask(
            page_items=right_pages,
            image_max_side=task.image_max_side,
            retries_left=task.retries_left,
        ),
    )


def _downscale_task(task: _GroqTask) -> _GroqTask | None:
    if task.image_max_side <= settings.groq_min_image_max_side:
        return None

    reduced = int(task.image_max_side * settings.groq_image_side_reduction_factor)
    if reduced >= task.image_max_side:
        reduced = task.image_max_side - 1
    if reduced < settings.groq_min_image_max_side:
        reduced = settings.groq_min_image_max_side
    if reduced >= task.image_max_side:
        return None

    return _GroqTask(
        page_items=task.page_items,
        image_max_side=reduced,
        retries_left=task.retries_left,
    )


def _format_page_numbers(page_numbers: list[int]) -> str:
    if not page_numbers:
        return "[]"
    if len(page_numbers) == 1:
        return str(page_numbers[0])
    return f"{page_numbers[0]}-{page_numbers[-1]}"


def _validate_batch_pages(batch_result: dict[str, Any], expected_pages: list[int]) -> None:
    pages = batch_result.get("pages", [])
    actual_page_numbers: list[int] = []
    for page in pages:
        raw_page_number = page.get("page_number")
        try:
            actual_page_numbers.append(int(raw_page_number))
        except (TypeError, ValueError) as exc:
            raise GroqConversionError(
                f"Groq batch output has invalid page number: {raw_page_number!r}"
            ) from exc

    if len(actual_page_numbers) != len(set(actual_page_numbers)):
        raise GroqConversionError(f"Groq batch output has duplicate pages: {actual_page_numbers}")

    expected_set = set(expected_pages)
    actual_set = set(actual_page_numbers)
    missing = sorted(expected_set - actual_set)
    extra = sorted(actual_set - expected_set)
    if missing or extra:
        raise GroqConversionError(
            "Groq batch output incomplete "
            f"(expected={sorted(expected_set)}, actual={sorted(actual_set)}, missing={missing}, extra={extra})"
        )


def _merge_batch_results(batch_results: list[dict[str, Any]]) -> dict[str, Any]:
    """Merge results from multiple batches into single document structure."""
    
    if not batch_results:
        raise GroqConversionError("No batch results to merge")
    
    # Use title from first batch
    merged = {
        "title": batch_results[0].get("title", ""),
        "pages": []
    }
    
    # Combine all pages from all batches
    for batch_result in batch_results:
        pages = batch_result.get("pages", [])
        merged["pages"].extend(pages)
    
    # Sort pages by page number to ensure correct order
    merged["pages"].sort(key=lambda p: p.get("page_number", 0))
    
    logger.info("Merged %d batches into %d total pages", len(batch_results), len(merged["pages"]))
    return merged


def _validate_merged_pages(content: dict[str, Any], expected_pages: int) -> None:
    """Ensure merged Groq result contains a complete contiguous page set."""
    pages = content.get("pages", [])
    page_numbers: list[int] = []
    for page in pages:
        raw_page_number = page.get("page_number")
        if raw_page_number is None:
            continue
        try:
            page_numbers.append(int(raw_page_number))
        except (TypeError, ValueError) as exc:
            raise GroqConversionError(
                f"Groq output has invalid page number: {raw_page_number!r}"
            ) from exc

    if len(page_numbers) != len(set(page_numbers)):
        raise GroqConversionError(f"Groq output has duplicate pages: {page_numbers}")

    expected = set(range(1, expected_pages + 1))
    actual = set(page_numbers)
    missing = sorted(expected - actual)
    extra = sorted(actual - expected)

    if missing or extra:
        logger.error(
            "Groq page validation failed. expected=%s actual=%s missing=%s extra=%s",
            sorted(expected),
            sorted(actual),
            missing,
            extra,
        )
        raise GroqConversionError(
            f"Groq output incomplete (missing pages: {missing}, extra pages: {extra})"
        )


def _create_docx_from_content(content: dict[str, Any], output_path: Path) -> None:
    """Create DOCX document from structured content with advanced formatting."""
    
    def safe_text(value: Any) -> str:
        """Safely convert any value to string, handling lists."""
        if isinstance(value, list):
            return ' '.join(str(item) for item in value if item)
        return str(value) if value else ""
    
    doc = Document()

    # Title comes from page sections — no separate heading to avoid duplication
    pages = content.get("pages", [])
    logger.info("Processing %d pages for DOCX creation", len(pages))
    
    for page_idx, page in enumerate(pages):
        page_num = page.get("page_number", page_idx + 1)
        sections = page.get("sections", [])
        logger.info("Page %d: processing %d sections", page_num, len(sections))
        
        for section_idx, section in enumerate(sections):
            section_type = section.get("type", "paragraph")
            text_content = safe_text(section.get("content", "")).strip()
            formatting = section.get("formatting", {})
            layout = section.get("layout", {})
            table_data = section.get("table_data", {})
            
            if not text_content and not table_data:
                if section_type == "table":
                    logger.warning("Skipping table section with empty table_data on page %d (section %d)", page_num, section_idx)
                else:
                    logger.debug("Skipping empty section %d on page %d", section_idx, page_num)
                continue
            
            logger.debug("Adding %s: %s", section_type, text_content[:100])
            
            # Handle different section types
            if section_type == "header_row":
                _add_header_row(doc, layout, formatting, text_content)
            elif section_type == "heading":
                _add_formatted_heading(doc, text_content, formatting)
            elif section_type == "table" and table_data:
                _add_formatted_table(doc, table_data, text_content)
            elif section_type == "list":
                _add_formatted_list(doc, text_content, formatting)
            else:  # paragraph
                _add_formatted_paragraph(doc, text_content, formatting)
    
    # Save document
    doc.save(output_path)
    logger.info("DOCX document saved to %s", output_path)


def _add_header_row(doc, layout: dict, formatting: dict, content: str = "") -> None:
    """Add header row with left and right aligned text.
    
    If layout.left_text and layout.right_text are both empty (common when Groq
    fills only the 'content' field), falls back to rendering 'content' as a
    heading paragraph so nothing is silently dropped.
    """
    left_text = layout.get("left_text", "")
    right_text = layout.get("right_text", "")
    
    if not left_text and not right_text:
        # Fallback: render the section content as a heading
        if content:
            font_size = formatting.get("font_size", "normal")
            level = 1 if font_size == "large" else 2
            para = doc.add_heading(content, level=level)
            _apply_paragraph_alignment(para, formatting)
            logger.debug("_add_header_row: rendered content as heading level=%d: %s", level, content[:80])
        return
        
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add left text
    if left_text:
        left_run = para.add_run(left_text)
        _apply_text_formatting(left_run, formatting)
    
    # Add tab to push right text to the right
    para.add_run("\t")
    
    # Add right text
    if right_text:
        right_run = para.add_run(right_text)
        _apply_text_formatting(right_run, formatting)
    
    # Set tab stop at right margin
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def _add_formatted_heading(doc, text: str, formatting: dict) -> None:
    """Add heading with formatting."""
    para = doc.add_heading(text, level=2)
    
    if para.runs:
        run = para.runs[0]
        _apply_text_formatting(run, formatting)
    
    _apply_paragraph_alignment(para, formatting)


def _add_formatted_paragraph(doc, text: str, formatting: dict) -> None:
    """Add paragraph with formatting.
    
    Splits text on newlines so that multi-line content from Groq
    (e.g. bulleted job-duty lists joined by \\n) becomes individual
    DOCX paragraphs instead of one compressed block.
    """
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        _add_text_with_number_highlighting(para, line, formatting)
        _apply_paragraph_alignment(para, formatting)


def _add_text_with_buyrug_coloring(para, text: str, base_formatting: dict) -> None:
    """Add text with special green coloring for BUYRUG words."""
    import re
    
    # Find BUYRUG-related words and make them green
    words = re.split(r'(\s+)', text)  # Split by whitespace but keep separators
    
    for word in words:
        if 'BUYRUG' in word.upper() or 'BUYRUQ' in word.upper():
            # Make BUYRUG words green and bold
            run = para.add_run(word)
            green_formatting = base_formatting.copy()
            green_formatting["color"] = "green"
            green_formatting["bold"] = True
            _apply_text_formatting(run, green_formatting)
        elif word.isdigit():
            # Keep number highlighting
            run = para.add_run(word)
            number_formatting = base_formatting.copy()
            number_formatting["bold"] = True
            _apply_text_formatting(run, number_formatting)
        else:
            # Regular text
            run = para.add_run(word)
            _apply_text_formatting(run, base_formatting)


def _add_formatted_list(doc, text: str, formatting: dict) -> None:
    """Add formatted list items."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines:
        para = doc.add_paragraph(line, style='List Bullet')
        if para.runs:
            _apply_text_formatting(para.runs[0], formatting)


def _add_formatted_table(doc, table_data: dict, fallback_text: str) -> None:
    """Add properly formatted table with borders and optimized column widths."""
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    has_borders = table_data.get("has_borders", True)
    header_styling = table_data.get("header_styling", True)
    column_types = table_data.get("column_types", [])
    
    if not headers and not rows:
        # Fallback to parsing table from text
        _add_table_from_text(doc, fallback_text)
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = str(header)
        
        if header_styling:
            # Bold header text
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            # Center align headers
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = str(cell_data)
                # Right align numbers
                if str(cell_data).replace(' ', '').replace(',', '').isdigit():
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Optimize column widths based on content
    _auto_adjust_column_widths(table, headers, column_types)
    
    # Apply table styling
    if has_borders:
        _set_table_borders(table)


def _add_table_from_text(doc, text: str) -> None:
    """Parse table from text and create formatted table with optimized widths."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    table_rows = []
    
    for line in lines:
        if '|' in line:
            # Parse pipe-separated table
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                table_rows.append(cells)
    
    if not table_rows:
        # Just add as paragraph if no table structure found
        para = doc.add_paragraph(text)
        return
    
    # Create table from parsed data
    max_cols = max(len(row) for row in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = table_rows[0] if table_rows else []
    
    for i, row_data in enumerate(table_rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell_data
                
                # Style first row as header
                if i == 0:
                    cell = table.rows[i].cells[j]
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Right align numbers in data rows
                elif str(cell_data).replace(' ', '').replace(',', '').replace('.', '').isdigit():
                    table.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Auto-adjust column widths
    _auto_adjust_column_widths(table, headers, [])
    
    _set_table_borders(table)


def _add_text_with_number_highlighting(para, text: str, base_formatting: dict) -> None:
    """Add text with numbers highlighted."""
    import re
    
    # Split text by numbers
    parts = re.split(r'(\d+)', text)
    
    for part in parts:
        if part.isdigit():
            # Highlight numbers
            run = para.add_run(part)
            number_formatting = base_formatting.copy()
            number_formatting["bold"] = True
            _apply_text_formatting(run, number_formatting)
        else:
            # Regular text
            run = para.add_run(part)
            _apply_text_formatting(run, base_formatting)


def _apply_text_formatting(run, formatting: dict) -> None:
    """Apply formatting to a text run."""
    if formatting.get("bold"):
        run.bold = True
    
    if formatting.get("italic"):
        run.italic = True
    
    # Apply color - expanded color support
    color = formatting.get("color", "black")
    if color == "green":
        run.font.color.rgb = RGBColor(0, 128, 0)  # Standard green
    elif color == "blue":
        run.font.color.rgb = RGBColor(0, 0, 255)
    elif color == "red":
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == "dark_blue":
        run.font.color.rgb = RGBColor(0, 0, 139)
    elif color == "orange":
        run.font.color.rgb = RGBColor(255, 165, 0)
    # Default is black, no need to set explicitly
    
    # Apply font size
    font_size = formatting.get("font_size", "normal")
    if font_size == "large":
        run.font.size = Pt(14)
    elif font_size == "small":
        run.font.size = Pt(10)
    else:
        run.font.size = Pt(12)


def _auto_adjust_column_widths(table, headers: list, column_types: list) -> None:
    """Automatically adjust column widths based on content type."""
    try:
        for i, header in enumerate(headers):
            header_lower = str(header).lower()
            
            # Determine column type
            if i < len(column_types):
                col_type = column_types[i]
            else:
                # Auto-detect from header name
                if header_lower in ['№', '#', 'no', 'num']:
                    col_type = "number"
                elif 'nomi' in header_lower or 'название' in header_lower or 'name' in header_lower:
                    col_type = "text"
                else:
                    col_type = "number"  # Default for numeric columns
            
            # Set width based on type
            if col_type == "number" and header_lower in ['№', '#', 'no', 'soni']:
                # Very narrow for row numbers and quantities
                table.columns[i].width = Inches(0.6)
            elif col_type == "text" or 'nomi' in header_lower:
                # Wide for text descriptions
                table.columns[i].width = Inches(3.2)
            elif header_lower in ['narxi', 'jami', 'price', 'total', 'сумма']:
                # Medium width for money amounts
                table.columns[i].width = Inches(1.4)
            else:
                # Default medium width
                table.columns[i].width = Inches(1.2)
                
    except Exception as e:
        logger.warning("Could not adjust column widths: %s", e)
        # Continue without width adjustment


def _apply_paragraph_alignment(para, formatting: dict) -> None:
    """Apply paragraph alignment."""
    alignment = formatting.get("alignment", "left")
    if alignment == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justified":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_table_borders(table) -> None:
    """Set borders for table."""
    try:
        # This is a simplified border setting - full implementation would need more XML manipulation
        for row in table.rows:
            for cell in row.cells:
                # Set cell borders through XML if needed
                pass
    except Exception as e:
        logger.warning("Could not set table borders: %s", e)
